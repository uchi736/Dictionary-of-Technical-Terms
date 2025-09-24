"""
Document loading utilities for various file formats
"""

import os
from pathlib import Path
from typing import Union, List, Optional
import logging

logger = logging.getLogger(__name__)


class DocumentLoader:
    """ドキュメントローダー - 各種フォーマットのファイル読み込み"""

    SUPPORTED_FORMATS = {
        '.txt': 'text',
        '.md': 'text',
        '.csv': 'csv',
        '.json': 'json',
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc'
    }

    @classmethod
    def load(cls, file_path: Union[str, Path]) -> str:
        """
        ファイルを読み込んでテキストを返す

        Args:
            file_path: ファイルパス

        Returns:
            読み込んだテキスト
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = file_path.suffix.lower()

        if ext not in cls.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {ext}")

        format_type = cls.SUPPORTED_FORMATS[ext]

        if format_type == 'text':
            return cls._load_text(file_path)
        elif format_type == 'csv':
            return cls._load_csv(file_path)
        elif format_type == 'json':
            return cls._load_json(file_path)
        elif format_type == 'pdf':
            return cls._load_pdf(file_path)
        elif format_type in ['docx', 'doc']:
            return cls._load_docx(file_path)
        else:
            raise ValueError(f"Handler not implemented for: {format_type}")

    @staticmethod
    def _load_text(file_path: Path) -> str:
        """テキストファイル読み込み"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def _load_csv(file_path: Path) -> str:
        """CSVファイル読み込み"""
        import csv
        text_parts = []

        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                text_parts.append(' '.join(row))

        return '\n'.join(text_parts)

    @staticmethod
    def _load_json(file_path: Path) -> str:
        """JSONファイル読み込み"""
        import json

        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # JSONを文字列化
        if isinstance(data, str):
            return data
        elif isinstance(data, list):
            return '\n'.join(str(item) for item in data)
        elif isinstance(data, dict):
            text_parts = []
            for key, value in data.items():
                text_parts.append(f"{key}: {value}")
            return '\n'.join(text_parts)
        else:
            return str(data)

    @staticmethod
    def _load_pdf(file_path: Path) -> str:
        """PDFファイル読み込み"""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(file_path))
            text_parts = []

            for page_num in range(doc.page_count):
                page = doc[page_num]
                text_parts.append(page.get_text())

            doc.close()
            return '\n'.join(text_parts)

        except ImportError:
            logger.warning("PyMuPDF not installed. Trying pdfplumber...")

            try:
                import pdfplumber
                text_parts = []

                with pdfplumber.open(str(file_path)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)

                return '\n'.join(text_parts)

            except ImportError:
                raise ImportError(
                    "PDF reading requires PyMuPDF or pdfplumber. "
                    "Install with: pip install pymupdf or pip install pdfplumber"
                )

    @staticmethod
    def _load_docx(file_path: Path) -> str:
        """Word文書読み込み"""
        try:
            from docx import Document
            doc = Document(str(file_path))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # テーブルからもテキストを抽出
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            return '\n'.join(text_parts)

        except ImportError:
            raise ImportError(
                "DOCX reading requires python-docx. "
                "Install with: pip install python-docx"
            )

    @classmethod
    def load_directory(
        cls,
        directory_path: Union[str, Path],
        pattern: str = "*",
        recursive: bool = False
    ) -> List[tuple]:
        """
        ディレクトリ内のファイルを読み込み

        Args:
            directory_path: ディレクトリパス
            pattern: ファイルパターン (例: "*.txt")
            recursive: サブディレクトリも検索するか

        Returns:
            (ファイルパス, テキスト)のリスト
        """
        directory_path = Path(directory_path)

        if not directory_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {directory_path}")

        results = []

        if recursive:
            files = directory_path.rglob(pattern)
        else:
            files = directory_path.glob(pattern)

        for file_path in files:
            if file_path.is_file() and file_path.suffix.lower() in cls.SUPPORTED_FORMATS:
                try:
                    text = cls.load(file_path)
                    results.append((str(file_path), text))
                    logger.info(f"Loaded: {file_path}")
                except Exception as e:
                    logger.error(f"Failed to load {file_path}: {e}")

        return results